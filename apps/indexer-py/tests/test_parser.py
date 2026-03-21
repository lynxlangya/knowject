from __future__ import annotations

import os
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document as DocxDocument
from openpyxl import Workbook
from pypdf import PdfWriter

from app.domain.indexing.parser import DocumentParser
from app.domain.indexing.pipeline import IndexDocumentRequest, IndexerError
from app.domain.indexing.runtime_config import IndexingRuntimeConfigResolver
from app.domain.indexing.embedding_client import EmbeddingRuntimeConfig
from app.domain.indexing.segments import ParsedSegment


def build_text_pdf_bytes(page_texts: list[str]) -> bytes:
    if not page_texts:
        raise ValueError("page_texts 不能为空")

    objects: dict[int, str] = {
        1: "<< /Type /Catalog /Pages 2 0 R >>",
    }
    page_object_ids: list[int] = []

    for index, page_text in enumerate(page_texts):
        page_object_id = 3 + index * 2
        content_object_id = page_object_id + 1
        page_object_ids.append(page_object_id)

        escaped_text = (
            page_text.replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )
        content = f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET"
        objects[content_object_id] = (
            f"<< /Length {len(content.encode('utf-8'))} >>\n"
            f"stream\n{content}\nendstream"
        )
        objects[page_object_id] = (
            "<< /Type /Page /Parent 2 0 R "
            "/MediaBox [0 0 612 792] "
            "/Resources << /Font << /F1 {font_id} 0 R >> >> "
            f"/Contents {content_object_id} 0 R >>"
        )

    font_object_id = 3 + len(page_texts) * 2
    objects[font_object_id] = "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    objects[2] = (
        "<< /Type /Pages "
        f"/Kids [{' '.join(f'{page_id} 0 R' for page_id in page_object_ids)}] "
        f"/Count {len(page_object_ids)} >>"
    )

    rendered_objects: dict[int, str] = {}
    for object_id, value in objects.items():
        rendered_objects[object_id] = value.format(font_id=font_object_id)

    chunks: list[bytes] = [b"%PDF-1.4\n"]
    offsets: dict[int, int] = {}

    for object_id in sorted(rendered_objects):
        offsets[object_id] = len(b"".join(chunks))
        chunks.append(
            f"{object_id} 0 obj\n{rendered_objects[object_id]}\nendobj\n".encode("utf-8")
        )

    start_xref = len(b"".join(chunks))
    max_object_id = max(rendered_objects)
    chunks.append(f"xref\n0 {max_object_id + 1}\n".encode("utf-8"))
    chunks.append(b"0000000000 65535 f \n")

    for object_id in range(1, max_object_id + 1):
        chunks.append(f"{offsets[object_id]:010d} 00000 n \n".encode("utf-8"))

    chunks.append(
        (
            f"trailer\n<< /Size {max_object_id + 1} /Root 1 0 R >>\n"
            f"startxref\n{start_xref}\n%%EOF\n"
        ).encode("utf-8")
    )
    return b"".join(chunks)


class ParserTest(unittest.TestCase):
    def setUp(self) -> None:
        resolver = IndexingRuntimeConfigResolver(error_factory=IndexerError)
        self.parser = DocumentParser(
            resolve_supported_extensions=resolver.resolve_supported_extensions,
            error_factory=IndexerError,
        )
        self.indexing_config = resolver.build_runtime_config()
        self.embedding_config = EmbeddingRuntimeConfig(
            provider="local_dev",
            api_key=None,
            base_url="https://api.openai.com/v1",
            model="text-embedding-3-small",
        )

    def test_parse_document_text_reads_utf8_markdown(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "demo.md"
            path.write_text("hello\n\nworld", encoding="utf-8")
            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="demo.md",
                    mime_type="text/markdown",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                text = self.parser.parse_document_text(request)

        self.assertEqual(text, "hello\n\nworld")

    def test_parse_document_segments_reads_markdown_as_single_segment(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "demo.md"
            path.write_text("hello\r\n\r\nworld", encoding="utf-8")
            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="demo.md",
                    mime_type="text/markdown",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                segments = self.parser.parse_document_segments(request)

        self.assertEqual(
            segments,
            [
                ParsedSegment(
                    text="hello\n\nworld",
                    source_kind="md",
                    order=0,
                )
            ],
        )

    def test_clean_text_collapses_extra_blank_runs(self):
        cleaned = self.parser.clean_text("a\r\n\r\n\r\nb\r\n\r\n\r\n\r\nc")

        self.assertEqual(cleaned, "a\n\n\nb\n\n\nc")

    def test_parse_document_segments_reads_digital_pdf_as_page_segments(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "demo.pdf"
            path.write_bytes(build_text_pdf_bytes(["Page one", "Page two"]))
            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="demo.pdf",
                    mime_type="application/pdf",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                segments = self.parser.parse_document_segments(request)

        self.assertEqual(
            segments,
            [
                ParsedSegment(
                    text="Page one",
                    source_kind="pdf",
                    order=0,
                    page_number=1,
                ),
                ParsedSegment(
                    text="Page two",
                    source_kind="pdf",
                    order=1,
                    page_number=2,
                ),
            ],
        )

    def test_parse_document_segments_rejects_pdf_without_extractable_text(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "blank.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=300)
            with path.open("wb") as file:
                writer.write(file)

            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="blank.pdf",
                    mime_type="application/pdf",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                with self.assertRaisesRegex(IndexerError, "OCR"):
                    self.parser.parse_document_segments(request)

    def test_parse_document_segments_reads_docx_with_heading_context(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "demo.docx"
            doc = DocxDocument()
            doc.add_heading("Overview", level=1)
            doc.add_paragraph("Intro paragraph")
            doc.add_heading("Details", level=2)
            doc.add_paragraph("Detail paragraph")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Key"
            table.cell(0, 1).text = "Value"
            doc.save(path)

            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="demo.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                segments = self.parser.parse_document_segments(request)

        self.assertEqual(
            segments,
            [
                ParsedSegment(
                    text="Overview\n\nIntro paragraph",
                    source_kind="docx",
                    order=0,
                    section_title="Overview",
                    heading_level=1,
                ),
                ParsedSegment(
                    text="Details\n\nDetail paragraph\n\nKey: Value",
                    source_kind="docx",
                    order=1,
                    section_title="Details",
                    heading_level=2,
                ),
            ],
        )

    def test_parse_document_segments_reads_xlsx_rows_with_header_path(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "demo.xlsx"
            workbook = Workbook()
            sheet1 = workbook.active
            sheet1.title = "Roadmap"
            sheet1.append(["Feature", "Status"])
            sheet1.append(["Upload", "Done"])
            sheet1.append(["Parser", "In Progress"])
            sheet2 = workbook.create_sheet("Metrics")
            sheet2.append(["Name", "Value"])
            sheet2.append(["Chunks", 42])
            workbook.save(path)

            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="demo.xlsx",
                    mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                segments = self.parser.parse_document_segments(request)

        self.assertEqual(
            segments,
            [
                ParsedSegment(
                    text="Roadmap | 第 2 行\nFeature: Upload\nStatus: Done",
                    source_kind="xlsx",
                    order=0,
                    sheet_name="Roadmap",
                    row_index=2,
                    header_path=("Feature", "Status"),
                ),
                ParsedSegment(
                    text="Roadmap | 第 3 行\nFeature: Parser\nStatus: In Progress",
                    source_kind="xlsx",
                    order=1,
                    sheet_name="Roadmap",
                    row_index=3,
                    header_path=("Feature", "Status"),
                ),
                ParsedSegment(
                    text="Metrics | 第 2 行\nName: Chunks\nValue: 42",
                    source_kind="xlsx",
                    order=2,
                    sheet_name="Metrics",
                    row_index=2,
                    header_path=("Name", "Value"),
                ),
            ],
        )

    def test_parse_document_segments_reads_single_row_xlsx_sheet_as_data(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "single-row.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Roadmap"
            sheet.append(["Upload", "Done"])
            workbook.save(path)

            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="single-row.xlsx",
                    mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                segments = self.parser.parse_document_segments(request)

        self.assertEqual(
            segments,
            [
                ParsedSegment(
                    text="Roadmap | 第 1 行\nColumn 1: Upload\nColumn 2: Done",
                    source_kind="xlsx",
                    order=0,
                    sheet_name="Roadmap",
                    row_index=1,
                    header_path=("Column 1", "Column 2"),
                )
            ],
        )

    def test_parse_document_segments_keeps_first_xlsx_row_when_header_row_is_not_text_like(self):
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "headerless.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Metrics"
            sheet.append([2026, 42])
            sheet.append([2027, 84])
            workbook.save(path)

            with patch.dict(os.environ, {"KNOWLEDGE_STORAGE_ROOT": temp_dir}):
                request = IndexDocumentRequest(
                    knowledge_id="knowledge-1",
                    document_id="document-1",
                    source_type="global_docs",
                    collection_name="global_docs",
                    file_name="headerless.xlsx",
                    mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    storage_path=str(path),
                    document_version_hash="hash-1",
                    embedding_config=self.embedding_config,
                    indexing_config=self.indexing_config,
                )

                segments = self.parser.parse_document_segments(request)

        self.assertEqual(
            segments,
            [
                ParsedSegment(
                    text="Metrics | 第 1 行\nColumn 1: 2026\nColumn 2: 42",
                    source_kind="xlsx",
                    order=0,
                    sheet_name="Metrics",
                    row_index=1,
                    header_path=("Column 1", "Column 2"),
                ),
                ParsedSegment(
                    text="Metrics | 第 2 行\nColumn 1: 2027\nColumn 2: 84",
                    source_kind="xlsx",
                    order=1,
                    sheet_name="Metrics",
                    row_index=2,
                    header_path=("Column 1", "Column 2"),
                ),
            ],
        )


if __name__ == "__main__":
    unittest.main()
